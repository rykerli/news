#!/usr/bin/env python
# -*- encoding: utf-8 -*-
"""
@Time    : 2019-01-11 13:05
@Author  : red
@Site    : 
@File    : docx_util.py
@Software: PyCharm
"""
import os
import shutil

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

global doc
global last_err
last_err = ""


def save(dst):
    global doc
    global last_err
    try:
        doc.save(dst)
        return True
    except IOError as e:
        last_err = e.strerror
        return False


def path_exists(path):
    if not os.path.exists(path):
        os.mkdir(path)
    else:
        shutil.rmtree(path)
        os.mkdir(path)


def add_info_title(title):
    global doc
    head = doc.add_heading("", level=3)
    run = head.add_run(title)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = u"宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


def add_text(text):
    global doc
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.name = u"宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


def add_text_list(text_list):
    for text in text_list:
        global doc
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(8)
        run.font.name = u"宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


def init(title):
    global doc
    doc = Document()
    doc.styles["Normal"].font.name = u'宋体'
    doc.styles["Normal"].font.size = Pt(8)
    doc.styles["Normal"]._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    head = doc.add_heading("", level=1)
    run = head.add_run(title)
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(128, 0, 0)
    run.font.name = u"宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def save_docx(array, path):
    init(array[0])
    add_info_title(u"标题")
    add_text(array[0])
    add_info_title(u"采集时间")
    add_text(array[1])
    add_info_title(u"正文")
    add_text(array[2])
    add_info_title(u"阅读数")
    add_text(array[3])
    add_info_title(u"评论")
    add_text_list(array[4])

    save(path)


def save_sina_docx(array, path):
    init(array[0])
    add_info_title(u"用户昵称")
    add_text(array[1])

    add_info_title(u"发布时间")
    add_text(str(array[2]))

    add_info_title(u"微博内容")
    add_text(array[3])

    add_info_title(u"转发数")
    add_text(str(array[4]))

    add_info_title(u"评论数")
    add_text(str(array[5]))

    add_info_title(u"点赞数")
    add_text(str(array[6]))

    add_info_title(u"评论")
    add_text_list(array[7])

    save(path)


if __name__ == '__main__':
    array = ['酒店推集赞吃免费大餐后反悔，遇上数百位政法大学生，结果……', '2018-11-18 21:20',
             '原标题：酒店推集赞吃免费大餐后反悔，遇上数百位政法大学生，结果……\n            近日，武汉一酒店通过官方微信发出一篇活动推文，称网友只要集齐80个点赞就可以免费领取一张价值168元的自助餐券。 \n11月15日，约400位学生前来兑换时，酒店却不予兑换。现场很多都是中南财经政法大学法学院的学生，积极维权。经警方和消协调解，酒店答应按承诺给券。 \n酒店承诺“集赞免费送餐券”却变卦 \n近日，有网友在微博爆料称，中南财经政法大学旁一酒店发布集赞送免费自助餐的活动，商家在发布该内容的文章下，回复“没有名额限制”。 \n \n当天，许多同学在朋友圈转发该文章求赞。然而，该商家突然通知“名额已满”。随后，许多学生前往酒店维权。 \n \n数百政法学子现场维权 \n15日下午1时许，记者在武汉民族大道金谷国际酒店内看到，近百名大学生正排成两队，等待领取自助餐券。一名学生表示，现在酒店没有工作人员发券，他们不知道能否领到餐券。 \n \n一名王姓同学说，14日她通过同学得知金谷国际酒店的官方微信发出一篇活动推广文章，文章称只要集齐80个点赞就可以免费领取酒店价值168元的西餐自助晚餐券一张，且没有名额限制，时间截止到11月16日中午12时。当日上午她和很多同学集齐80个点赞后到金谷国际酒店来领取自助餐券时被告知，该活动是针对曾在酒店消费过的客人，不针对学生。 \n \n还有人称，服务员现场有恃无恐“有本事你举报我啊，你们又没消费。” \n \n酒店的说法引发现场学子的愤怒，大伙围住酒店工作人员欲讨个说法。学生报警后，警察叔叔也及时赶到，并组织学生有秩序地登记相关情况。 \n在警方和消协部门的协调下，酒店方最终向所有集齐80个点赞的一千多名学生发放了自助餐券。 \n \n视频曝光，维权被网友怒赞 \n这次事件本就这样过去了，但是当网上曝出政法学子现场维权视频时，又引发了新一波的关注。 \n有网友评价该视频：“内容引起强烈舒适。” \n对于为什么要维权，中南财经政法大学的学生是这样说的： \n“金谷酒店在微信公众号里发起活动，说集满80个赞，就可以免费获得168元的自助餐券一张。然后我们学校的很多学生都集赞了。他们今天早上的时候突然一下子不承认这个活动了，说名额已经满了。我们在那个推文下面，已经明确地问过他们，是否有名额限制这一条，他们回复了两次，是没有名额限制的。他们那边相当于给我们发出了一个要约，而我们这边也承诺了，那我们这边按照合同法，我们就应该有这样一份168的晚餐。” \n是不是有理有利有节，让人心服口服？ \n当天，该视频还登上微博、知乎等平台热搜，@人民日报 19时发布了相关内容的微博：学以致用！酒店举办“集赞免费吃”又反悔，400名政法大学生维权]15日，@中南财经政法大学 旁一酒店发布集赞换代金券免费吃大餐活动，约400位学生前来兑换时，酒店却不予兑换。现场很多都是中南财大法学院的学生，积极维权。经警方和消协调解，酒店答应按承诺给券。 \n \n@中南财经政法大学 还转发人民日报微博称：“就地证法”。 \n \n网友评论 \n \n于维权，网友也提出了非常正确的观点：维权过程中要保持良好的秩序，只用事实和法律讲话。同时，不影响酒店的其他客人，不影响警察执法，依法维权。 \n来源：楚天都市报、澎湃新闻返回搜狐，查看更多      责任编辑：',
             None, ['胡蘿蔔素過多症(北京市北京市)', '无良商家玩套路，政法学子讲证据😄😄😄😄', None, None]]
    save_docx(array, os.path.join("/Users/red/Desktop", array[0] + ".docx"))
